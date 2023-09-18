# from materiais.funcs import replace_in_runs
import re
import docx

PATTERN_SUP_TAG = re.compile(r"(\d+(\.\d+)?x10\s*)(\-?\d+)")
PATTERN_CONTEUDO = re.compile(r"(.+?) / ?(.+?)\n")
PATTERN_ALTERNATIVAS = re.compile(r"\b(a|b|c|d|e)\)\s")
REPLACEMENT_ALTERNATIVAS = r"\1$ "
PATTERN_GAB = re.compile(r"Gab.*?([a-eA-E])", re.IGNORECASE)
REPLACEMENT_GAB = r"w$ \1"


# Remove when done
def docx_to_text(doc_obj):
    """
    Retorna o texto inteiro do .docx
    """
    full_text = []
    for paragraph in doc_obj.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


# Remove when done
def replace_in_runs(paragraph, pattern, replacement):
    for run in paragraph.runs:
        if pattern.search(run.text):
            print(f"Encontrado em {run.text}")
            run.text = pattern.sub(replacement, run.text)


'''def trata_conteudos(texto_extraido_do_docx_em_string: str) -> str:
    """
    Aplica conteudos no texto e retorna o texto em str com as modificacoes.
    """
    texto_modificado = texto_extraido_do_docx_em_string
    matches = list(PATTERN_CONTEUDO.finditer(texto_extraido_do_docx_em_string))
    for match in matches:
        whole, conteudo, sub_conteudo = match.group(0), match.group(1), match.group(2)
        if len(conteudo) < 50 and len(sub_conteudo) < 50:
            replace = f"Conteudo: {conteudo} / {sub_conteudo};\n"
            texto_modificado = texto_modificado.replace(whole, replace, 1)
    return texto_modificado'''


def apply_sup_tags(texto_extraido_do_docx_em_string: str) -> str:
    """
    Aplica suptags em notacoes cientificas, e retorna o texto completo aplicado.
    """
    texto_modificado = texto_extraido_do_docx_em_string
    matches = list(PATTERN_SUP_TAG.finditer(texto_extraido_do_docx_em_string))
    for match in matches:
        whole, base, expoente = match.group(0), match.group(1), match.group(3)
        if "<sup>" not in whole:
            texto_modificado = texto_modificado.replace(
                whole, f"{base}<sup>{expoente}</sup>"
            )

    return texto_modificado


def trata_alternativas(doc_obj, doc_path) -> None:
    """
    Salva arquivo com alternativas tratadas para extracao.
    """
    for para in doc_obj.paragraphs:
        replace_in_runs(para, PATTERN_ALTERNATIVAS, REPLACEMENT_ALTERNATIVAS)
    doc_obj.save(doc_path)


def replace_gab_runs_divided(paragraph):
    runs = paragraph.runs
    for i in range(len(runs) - 1):
        combined_text = runs[i].text + runs[i + 1].text
        if "Gab:" in combined_text:
            runs[i].text = combined_text.replace("Gab:", "w$")
            runs[i + 1].text = ""


def trata_gabs(doc_obj, doc_path):
    for para in doc_obj.paragraphs:
        replace_gab_runs_divided(para)
    doc_obj.save(doc_path)


def possui_img_tag(text: str) -> bool:
    return "[IMG]" in text


class tratamento_geral_pra_extracao:
    def __init__(self, doc_obj, doc_path) -> None:
        self.doc_obj = doc_obj
        self.doc_path = doc_path

    def Tratamento(self):
        trata_gabs(self.doc_obj, self.doc_path)
        trata_alternativas(self.doc_obj, self.doc_path)

    def Tratamento_sup_tags(self, texto):
        return apply_sup_tags(texto)

    def checa_imagens_questoes(self) -> list:
        """'
        Checa se tem imagem em cada parte da questao.\n
        \t questao = {\n
        \t    "imagem_no_enunciado": False/True,\n
        \t    "imagem_na_a": False/True,\n
        \t    "imagem_na_b": False/True,\n
        \t    "imagem_na_c": False/True,\n
        \t    "imagem_na_d": False/True,\n
        \t    "imagem_na_e": False/True,\n
        }
        """
        estamos_no_enunciado = False
        questoes = []
        questao = {
            "imagem_no_enunciado": False,
            "imagem_na_a": False,
            "imagem_na_b": False,
            "imagem_na_c": False,
            "imagem_na_d": False,
            "imagem_na_e": False,
        }
        for para in self.doc_obj.paragraphs:
            texto = para.text
            if "Questão-" in texto:
                if questao:
                    questoes.append(questao)
                questao = {
                    "imagem_no_enunciado": False,
                    "imagem_na_a": False,
                    "imagem_na_b": False,
                    "imagem_na_c": False,
                    "imagem_na_d": False,
                    "imagem_na_e": False,
                }
                estamos_no_enunciado = True

            if re.match(r"[a-e]\$", texto):
                estamos_no_enunciado = False

            if estamos_no_enunciado:
                if possui_img_tag(texto):
                    questao["imagem_no_enunciado"] = True

            if not estamos_no_enunciado:
                if "a$" in texto and possui_img_tag(texto):
                    questao["imagem_na_a"] = True
                if "b$" in texto and possui_img_tag(texto):
                    questao["imagem_na_b"] = True
                if "c$" in texto and possui_img_tag(texto):
                    questao["imagem_na_c"] = True
                if "d$" in texto and possui_img_tag(texto):
                    questao["imagem_na_d"] = True
                if "e$" in texto and possui_img_tag(texto):
                    questao["imagem_na_e"] = True
        if questao:
            questoes.append(questao)
        questoes.pop(0)
        return questoes

    def trata_conteudos(self, texto_extraido_do_docx_em_string: str) -> str:
        """
        Aplica conteudos no texto e retorna o texto em str com as modificacoes.
        """
        texto_modificado = texto_extraido_do_docx_em_string
        matches = list(PATTERN_CONTEUDO.finditer(texto_extraido_do_docx_em_string))
        for match in matches:
            whole, conteudo, sub_conteudo = (
                match.group(0),
                match.group(1),
                match.group(2),
            )
            if len(conteudo) < 50 and len(sub_conteudo) < 50:
                replace = f"Conteudo: {conteudo} / {sub_conteudo};\n"
                texto_modificado = texto_modificado.replace(whole, replace, 1)
        return texto_modificado
