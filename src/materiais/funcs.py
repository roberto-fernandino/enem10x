from pathlib import Path
import os
from django.utils.text import slugify
import docx
from docx2txt import process
from docx import Document
import os

BASE_DIR = Path(__file__).resolve().parent.parent



def alternativas(arquivo):
    doc = docx.Document(arquivo)
    texto_completo = "\n".join(
        [paragrafo.text for paragrafo in doc.paragraphs]
    )  # Todo paragrafo existente no documento sendo adicionado a variavel
    questoes = texto_completo.split("Questão")[
        1:
    ]  # Dividindo todos paragrafos por questão, Assim salvando apenas a Questão, enunciado e alternativas
    alternativas_global = (
        []
    )  # Lista de dicionarios, com alternativas de cada questão, um dicionario por questão

    for num, questao in enumerate(questoes, 1):  # Começar a enumerar de 1
        alternativas_questao = {}  # Dicionario com as alternativas para cada questão
        # Buscando as alternativas
        for index, letra in enumerate(
            ["a)", "b)", "c)", "d)", "e)", 'w$', 'm$', 's$', 'n$','t$', 'p$']
        ):  # Citando todas as alternativas possiveis , no caso 5
            inicio = questao.find(letra)
            # Se está na ultima letra (e) definimos que é a ultima alternativa
            if index == 10:
                fim = len(questao)  # Definição do Fim
            else:
                # Procurando o início da próxima alternativa para definir o fim da alternativa atual
                proxima_letra = ["a)", "b)", "c)", "d)", "e)", 'w$', 'm$', 's$', 'n$','t$', 'p$'][index + 1]
                fim = questao.find(proxima_letra)

            if (
                inicio != -1 and fim != -1
            ):  # Certifica-se de que a variável inicio é diferente de -1 e Certifica-se de que a variável fim é diferente de -1.
                alternativas_questao[letra[0]] = questao[
                    inicio + 2 : fim
                ].strip()  # Atribui a proxima letra se não for a ultima
        alternativas_global.append(
            alternativas_questao
        )  # Adionando a lista o dicionario contido com as alternativas de cada questão
            

    return alternativas_global  # Retorna alternativa_global
#print(alternativas('/home/roberto/projects/enem10x/src/leitores/questions3.docx'))
def extrair_enunciados(arquivo):
    lista_enunciados = []

    doc = docx.Document(arquivo)
    texto_completo = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
    questoes = texto_completo.split("Questão")[1:]

    for questao in questoes:
        # Encontrar o início e o fim do enunciado da questão
        inicio_enunciado = questao.find("-") + 1
        fim_enunciado = min(
            [
                questao.find(letra)
                for letra in ["a)", "b)", "c)", "d)", "e)"]
                if questao.find(letra) != -1
            ]
        )  # Procurando as alternativas e quando achar volta um argumento
        # O minimo da lista é a primeira alternativa de resposta ou seja quando começa as alternativas, acaba o enunciado
        # Extrair o enunciado e adicionar à lista
        enunciado = questao[inicio_enunciado:fim_enunciado].strip()
        lista_enunciados.append(enunciado)

    return lista_enunciados

#print(extrair_enunciados('/home/roberto/projects/enem10x/src/leitores/questions3.docx'))

def define_image_path(instance, filename: str) -> str:
    ext = os.path.splitext(filename)[1]
    filename = f'enem.{ext.lstrip(".")}'
    return Path(f"questoes/{filename}png")


def define_ranking_conteudo_prova(conteudos_errados: list, conteudos_acertados: list):
    """Gera um ranking de conteudos mais errados, acertados respectivamente de uma prova especifica e os retorna no formato:
    tuple(ranking errados, ranking acertados)"""

    ranking = {}
    maior_contagem = 0
    for conteudo in conteudos_errados:
        ranking[conteudo] = ranking.get(conteudo, 0) + 1
        if ranking[conteudo] > maior_contagem:
            maior_contagem = ranking[conteudo]
    ranking_erradas = dict(
        sorted(ranking.items(), key=lambda item: item[1], reverse=True)
    )

    ranking = {}
    maior_contagem = 0
    for conteudo in conteudos_acertados:
        ranking[conteudo] = ranking.get(conteudo, 0) + 1
        if ranking[conteudo] > maior_contagem:
            maior_contagem = ranking[conteudo]

    ranking_certos = dict(
        sorted(ranking.items(), key=lambda item: item[1], reverse=True)
    )

    return (
        ranking_erradas,
        ranking_certos,
    )


# Inutil agora
def retorna_tipos_prova(tipos_prova_tratado: list) -> str:
    tipos = "  -  ".join(tipos_prova_tratado)
    return tipos

def procura_imagens() -> list:
        imagens = os.listdir(f"{BASE_DIR}/media/imagens")
        return imagens

